#!/usr/bin/env python3
"""
Government Document -> Audit Checklist Generator
Uses BitNet (bitnet.cpp / llama-cli) for local offline inference.

Usage (standalone CLI):
    python3 checklist_generator.py your_document.pdf [max_clauses]

Environment variables (override defaults):
    BITNET_CLI    — path to the llama-cli binary
    BITNET_MODEL  — path to the .gguf model file
"""
import os, sys, re, json, math, subprocess, tempfile

PDF_PATH    = "document.pdf"
OUTPUT_DOCX = "Audit_Checklist.docx"
MAX_ITEMS   = 10

# ── Configure these via environment variables ────────────────────────────────
# See README.md for setup instructions.
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def _find_model():
    # 1. Check environment variable first
    env_model = os.environ.get("BITNET_MODEL")
    if env_model and os.path.exists(env_model):
        return env_model

    # 2. Check models/ folder
    models_dir = os.path.join(BASE_DIR, "models")
    if not os.path.isdir(models_dir):
        try: os.makedirs(models_dir)
        except: pass
        return None

    for f in sorted(os.listdir(models_dir)):
        if f.endswith(".gguf"):
            return os.path.join(models_dir, f)
    return None

MODEL_PATH = _find_model() or os.environ.get("BITNET_MODEL", "")

if len(sys.argv) >= 2: PDF_PATH  = sys.argv[1]
if len(sys.argv) >= 3: MAX_ITEMS = int(sys.argv[2])

DEPTH_PRESETS = {
    "light":    {"n": 1, "m": 25},
    "standard": {"n": 2, "m": 50},
    "deep":     {"n": 3, "m": 75},
    "full":     {"n": 9999, "m": 100},
}

# ── Clause extraction patterns ───────────────────────────
DOTTED_RE = re.compile(
    r'^((?:\d+|[A-Z])(?:\.[A-Z0-9]+){1,5})\s{1,6}(\S.{8,})',
    re.IGNORECASE
)
NUMBERED_RE = re.compile(r'^(\d{1,2})\.\s{1,4}(\S.{3,})')
LETTERED_RE = re.compile(r'^\(([a-z])\)\s{1,4}(\S.{8,})')
CLAUSE_RE   = DOTTED_RE

SECTION_RE_FALLBACK = re.compile(
    r'^(PART|SECTION|CHAPTER|SUBPART|Subpart|ANNEX|APPENDIX)\s*[-–]?\s*([A-Z0-9]+)\b',
    re.IGNORECASE
)
HEADER_RE = re.compile(
    r'^(INTRODUCTION|DEFINITIONS|CONCLUSION|PREFACE|QA\s+ACTIVITIES'
    r'|SAFETY|HANDLING|ASSEMBLY|ANNEXURE|APPENDIX)',
    re.IGNORECASE
)
OBLIGATION = re.compile(
    r'\b(shall|must|will\s+be|is\s+required|are\s+required|required\s+to|'
    r'to\s+be\s+(?:used|carried|checked|verified|performed|filled|submitted|'
    r'maintained|connected|fitted|followed|generated|screened|formed|accepted)|'
    r'ensure|maintain|establish|implement|document|verify|review|'
    r'check|inspect|audit|control|comply|responsible|'
    r'authoris|authoriz|monitor|record|report|assess|evaluate|screen)\b',
    re.IGNORECASE
)
SKIP_SECTIONS = {
    'DEFINITIONS','INTRODUCTION','PREFACE','CONCLUSION','BACKGROUND',
    'FOREWORD','SCOPE','PURPOSE','REFERENCES','GLOSSARY','ABBREVIATIONS',
    'ANNEXURE','ANNEX','APPENDIX','TABLEOFCONTENTS','TOC',
}

def _section_is_skippable(section_name):
    normalised = re.sub(r'[^A-Z]', '', section_name.upper())
    for skip in SKIP_SECTIONS:
        skip_norm = re.sub(r'[^A-Z]', '', skip.upper())
        if skip_norm and (skip_norm in normalised or normalised.startswith(skip_norm)):
            return True
    return False

SKIP_LINE = re.compile(
    r'^(Figure|Fig\.|Table\s+\d|Appendix|Annex|Note:|NOTE:|'
    r'Page\s+\d+|Page\s+\|\s+\d+|\d+$|[ivxlIVXL]+$|[-–—]{3,}|www\.|http)',
    re.IGNORECASE
)

# ── BitNet inference ─────────────────────────────────────
def _run_local_llm(llm, prompt, max_tokens=256, temperature=0.1):
    """
    Call the local Llama instance with Llama-3 chat format.
    """
    if llm is None:
        # Fallback for CLI usage if no LLM is passed
        return ""

    chat_prompt = (
        "<|begin_of_text|>"
        "<|start_header_id|>system<|end_header_id|>\n\n"
        "You are a precise government quality auditor assistant. "
        "Follow instructions exactly and output only what is asked.<|eot_id|>"
        "<|start_header_id|>user<|end_header_id|>\n\n"
        f"{prompt}<|eot_id|>"
        "<|start_header_id|>assistant<|end_header_id|>\n\n"
    )
    
    try:
        output = llm(
            chat_prompt,
            max_tokens=max_tokens,
            temperature=temperature,
            stop=["<|eot_id|>", "<|end_of_text|>"],
            echo=False
        )
        return output['choices'][0]['text'].strip()
    except Exception as e:
        print(f"Inference error: {e}")
        return ""

# ── Step 1: Read PDF ─────────────────────────────────────
def load_model(model_path):
    print(f"\n[4/5] Initializing engine (Auto-Detecting Hardware)...")
    if not model_path or not os.path.exists(model_path):
        sys.exit(f"ERROR: Model file not found at {model_path}")
    
    from llama_cpp import Llama
    import llama_cpp
    
    try:
        llm = Llama(
            model_path=model_path,
            n_ctx=8192,
            type_k=llama_cpp.GGML_TYPE_Q8_0,
            type_v=llama_cpp.GGML_TYPE_Q8_0,
            flash_attn=True,
            n_threads=os.cpu_count() or 4,
            n_gpu_layers=-1,
            verbose=False
        )
        print(f"   Success: TurboQuant Mode Enabled (8-bit KV + 8k Context).")
        return llm
    except Exception as e:
        print(f"   Hardware/Build check failed: {e}")
        print(f"   Switching to Universal Safe Mode (F16 KV + 4k Context).")
        return Llama(
            model_path=model_path, 
            n_ctx=4096, 
            n_gpu_layers=0, 
            verbose=False
        )

def read_pdf(path):
    print(f"\n[1/5] Reading: {path}")
    try:
        import pdfplumber
    except ImportError:
        sys.exit("ERROR: pip install pdfplumber")
    if not os.path.exists(path):
        sys.exit(f"ERROR: File not found - {path}")
    pages = []; scanned = 0
    with pdfplumber.open(path) as pdf:
        for pg in pdf.pages:
            try:
                t = pg.extract_text()
                text = t.strip() if t else ""
                if len(text) < 30:
                    scanned += 1; pages.append("")
                else:
                    pages.append(text)
            except Exception:
                scanned += 1; pages.append("")
    readable = sum(1 for p in pages if p)
    print(f"   {len(pages)} pages total — {readable} machine-readable, {scanned} scanned/empty (skipped).")
    if readable == 0:
        sys.exit("ERROR: No readable text. PDF appears fully scanned.")
    return pages

# ── Step 2: Detect structure ─────────────────────────────
STRUCTURE_PROMPT = (
    "Analyze this government regulatory document sample. "
    "Identify how major sections are labeled and give one example clause number.\n\n"
    "Respond with ONLY this JSON, no explanation:\n"
    '{{\n'
    '  "section_pattern": "^(Subpart\\\\s+[A-Z0-9]+)",\n'
    '  "section_example": "Subpart B1",\n'
    '  "clause_example": "21.B1.4"\n'
    '}}\n\n'
    "SAMPLE PAGES:\n{sample}"
)

def detect_structure(llm, pages):
    print("\n[2/5] Detecting document structure via LLM...")
    total   = len(pages)
    indices = sorted(set([*range(5, min(11, total)), *range(min(30, total), min(36, total))]))
    sample  = ""
    for i in indices:
        if i < total and pages[i].strip():
            sample += f"\n--- PAGE {i+1} ---\n{pages[i][:250]}\n"
        if len(sample) > 2000:
            break
    raw = _run_local_llm(llm, STRUCTURE_PROMPT.format(sample=sample), max_tokens=150, temperature=0.1)
    raw = re.sub(r"```json|```", "", raw).strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        try:
            s = json.loads(m.group())
            if "section_pattern" in s:
                print(f"   Section : {s.get('section_example','?')} -> {s.get('section_pattern')}")
                print(f"   Clause  : {s.get('clause_example','?')}")
                return s
        except Exception:
            pass
    print("   Could not detect structure — using fallback.")
    return None

# ── Step 3: Compile patterns ─────────────────────────────
def compile_patterns(structure):
    if not structure:
        return CLAUSE_RE, SECTION_RE_FALLBACK
    try:
        section_re = re.compile(structure["section_pattern"], re.IGNORECASE)
        assert section_re.groups >= 1
        return CLAUSE_RE, section_re
    except Exception:
        return CLAUSE_RE, SECTION_RE_FALLBACK

# ── Step 4: Extract clauses ──────────────────────────────
def clean(t):
    t = re.sub(r'(?<=[A-Za-z]) (?=[A-Za-z] )', '', t)
    return re.sub(r'\s{2,}', ' ', t).strip()

def _compile_section_filter(section_filter_str):
    """
    Parse a user-supplied section filter string into a list of regex patterns.
    Accepts comma/semicolon separated values like: "21 C1, 21.B, Subpart-A"
    Each token is turned into a case-insensitive prefix regex against clause ref.
    Returns None if no filter specified.
    """
    if not section_filter_str or not section_filter_str.strip():
        return None
    tokens = re.split(r'[,;]+', section_filter_str.strip())
    patterns = []
    for tok in tokens:
        tok = tok.strip()
        if not tok:
            continue
        escaped = re.escape(tok).replace(r'\*', '.*')
        patterns.append(re.compile(escaped, re.IGNORECASE))
    return patterns if patterns else None

def _clause_matches_filter(ref, filter_patterns):
    """Return True if ref matches any of the filter patterns (substring match)."""
    if filter_patterns is None:
        return True
    for pat in filter_patterns:
        if pat.search(ref):
            return True
    return False

def _is_page_header(line, current_section):
    if len(line) > 65:
        return True
    m = re.match(r'^(PART|SECTION|CHAPTER|SUBPART|Subpart|ANNEX|APPENDIX)\s*[-–]?\s*([A-Z0-9]+)\b',
                 line, re.IGNORECASE)
    if not m:
        return False
    candidate = re.sub(r'[^A-Za-z0-9\-]', '-',
                       re.sub(r'\s+', '-', m.group(0).strip())).strip('-')
    if candidate.upper() == current_section.upper():
        return True
    return False

def _is_inline_crossref(line):
    if re.search(r'\.\s*$', line):
        return True
    if re.search(r'\b(of|in|as per|under|per)\b', line, re.IGNORECASE):
        return True
    return False

def extract_clauses(pages, clause_re, section_re, section_filter=None):
    filter_patterns = _compile_section_filter(section_filter)
    print("\n[3/5] Extracting clauses...")
    if filter_patterns:
        print(f"   Section filter active: {section_filter!r}")

    current_section = ""; pending_num = None
    pending_section = ""; pending_body = []; clauses = []
    last_section_line = ""

    def flush():
        if pending_num and pending_body:
            if _section_is_skippable(pending_section): return
            body = clean(" ".join(pending_body))
            if len(body) > 20 and OBLIGATION.search(body):
                ref = f"{pending_section}/{pending_num}" if pending_section else pending_num
                if _clause_matches_filter(ref, filter_patterns) or \
                   _clause_matches_filter(pending_num, filter_patterns):
                    clauses.append({"ref": ref, "text": body})

    last_int_num = None
    for page_text in pages:
        lines = [l.strip() for l in page_text.split("\n")]
        i = 0
        while i < len(lines):
            line = lines[i]; i += 1
            if not line or SKIP_LINE.match(line): continue

            is_section = False
            for sr in [section_re, HEADER_RE]:
                sm = sr.match(line)
                if not sm:
                    continue
                if _is_page_header(line, current_section):
                    break
                if pending_num and _is_inline_crossref(line):
                    pending_body.append(line)
                    is_section = True
                    break
                label = sm.group(0).strip()[:45]
                candidate = re.sub(r'[^A-Za-z0-9\-]', '-',
                                   re.sub(r'\s+', '-', label)).strip('-')
                if candidate.upper() == current_section.upper():
                    is_section = True; break
                flush()
                current_section = candidate
                pending_num = None; pending_body = []
                last_int_num = None; is_section = True; break
            if is_section: continue

            dm = DOTTED_RE.match(line)
            if dm:
                clause_num = dm.group(1)
                if re.match(r'^\d+\.0$', clause_num):
                    flush(); pending_num = None; pending_body = []; continue
                title_parts = [dm.group(2)]
                while len(" ".join(title_parts)) < 10 and i < len(lines):
                    peek = lines[i].strip()
                    if not peek:
                        i += 1; continue
                    if (DOTTED_RE.match(peek) or NUMBERED_RE.match(peek) or
                            section_re.match(peek) or HEADER_RE.match(peek) or
                            SKIP_LINE.match(peek)):
                        break
                    if re.match(r'^(Regulation|Acceptable Means|Guidance Material|AMC|GM)\b',
                                peek, re.IGNORECASE):
                        break
                    title_parts.append(peek); i += 1
                flush()
                pending_num = clause_num; pending_section = current_section
                pending_body = title_parts[:]; last_int_num = None
                continue

            nm = NUMBERED_RE.match(line)
            if nm:
                body_text = nm.group(2).strip()
                if len(body_text) < 5 or re.match(r'.+\s+\d+\s+\d+$', body_text): continue
                flush()
                int_num = nm.group(1); pending_num = int_num
                pending_section = current_section; pending_body = [body_text]
                last_int_num = int_num; continue

            lm = LETTERED_RE.match(line)
            if lm:
                flush()
                letter = lm.group(1); body_text = lm.group(2).strip()
                parent = last_int_num if last_int_num else (pending_num or "")
                pending_num = f"{parent}({letter})" if parent else f"({letter})"
                pending_section = current_section; pending_body = [body_text]; continue

            if pending_num:
                if (re.match(r'^[A-Z][A-Z\s\-\/\.]{8,}$', line) and
                        len(line) < 60 and
                        not re.match(r'^(NIL|TBD|N\/A|NA)\s*$', line, re.IGNORECASE)):
                    flush(); pending_num = None; pending_body = []
                else:
                    pending_body.append(line)

    flush()
    seen, unique = set(), []
    for c in clauses:
        if c["ref"] not in seen:
            seen.add(c["ref"]); unique.append(c)
    print(f"   {len(unique)} unique obligation clauses found.")
    return unique



# ── Step 6: Extract checkpoints ──────────────────────────
DEFAULT_CHECKPOINT_PROMPT = (
    "Read the following clause from a government regulatory document.\n"
    "List every auditable requirement as a numbered YES/NO question (max 15 words each).\n"
    "Output ONLY the numbered list. No intro, no explanation.{focus}\n\n"
    "CLAUSE:\n{text}"
)

def extract_checkpoints(llm, clause, prompt_template=None, focus=""):
    template  = prompt_template if prompt_template else DEFAULT_CHECKPOINT_PROMPT
    focus_str = f"\nAuditor focus: {focus.strip()}" if focus and focus.strip() else ""
    try:
        filled = template.format(text=clause["text"][:800], focus=focus_str)
    except KeyError:
        filled = template.replace("{text}", clause["text"][:800])

    raw = _run_local_llm(llm, filled, max_tokens=400, temperature=0.15)

    checkpoints = []
    for line in raw.split("\n"):
        line = line.strip()
        if not line: continue
        cleaned = re.sub(r'^[\d]+[\.\)]\s*', '', line).strip()
        cleaned = re.sub(r'^[-•*]\s*', '', cleaned).strip()
        cleaned = cleaned.strip('"\'')
        cleaned = re.sub(r'\.+\?', '?', cleaned)
        cleaned = re.sub(r'\?+', '?', cleaned).strip()
        if len(cleaned) > 12 and not cleaned.lower().startswith(("clause", "here", "below", "note")):
            if not cleaned.endswith("?"):
                cleaned += "?"
            checkpoints.append(cleaned)

def extract_checkpoints_batch(llm, clauses, focus=""):
    """
    Turbo Mode: Extract checkpoints for multiple clauses in one go.
    """
    if not clauses: return []
    
    batch_text = ""
    for i, c in enumerate(clauses, 1):
        batch_text += f"CLAUSE {i} ({c['ref']}):\n{c['text'][:500]}\n\n"
        
    prompt = (
        "Read these regulatory clauses. For EACH clause, list the auditable YES/NO questions.\n"
        "Format your response as:\n"
        "CLAUSE 1: [Questions...]\n"
        "CLAUSE 2: [Questions...]\n"
        f"{focus}\n\n"
        f"{batch_text}"
    )
    
    raw = _run_local_llm(llm, prompt, max_tokens=2048, temperature=0.1)
    
    # Simple parsing: find which questions belong to which clause
    # (In a real scenario, this would be more robust, but this matches the 'Turbo' spirit)
    all_checkpoints = []
    current_clause_idx = 0
    
    lines = raw.split("\n")
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Check if we hit a new clause header in the output
        if re.search(r'CLAUSE\s+\d+', line, re.IGNORECASE):
            m = re.search(r'\d+', line)
            if m:
                current_clause_idx = int(m.group()) - 1
            continue
            
        if 0 <= current_clause_idx < len(clauses):
            cleaned = re.sub(r'^[\d]+[\.\)]\s*', '', line).strip()
            cleaned = re.sub(r'^[-•*]\s*', '', cleaned).strip()
            if len(cleaned) > 10:
                if not cleaned.endswith("?"): cleaned += "?"
                all_checkpoints.append({
                    "clause_ref": clauses[current_clause_idx]["ref"],
                    "question": cleaned
                })
                
    return all_checkpoints

def build_master_turbo(llm, clauses, max_items=0, focus="", batch_size=3):
    subset = clauses[:max_items] if max_items else clauses
    total  = len(subset)
    master = []
    
    for i in range(0, total, batch_size):
        batch = subset[i : i + batch_size]
        print(f"   [Turbo] Processing batch {i//batch_size + 1} ({len(batch)} clauses)...")
        
        batch_results = extract_checkpoints_batch(llm, batch, focus=focus)
        
        # Reformat into master structure
        for res in batch_results:
            master.append({
                "clause_base": res["clause_ref"],
                "subpoint": "T", # T for Turbo
                "ref": f"{res['clause_ref']} (T)",
                "question": res["question"]
            })
            
    return master

# ── Step 7: Filter checkpoints ───────────────────────────
def filter_checkpoints(master_rows, n, m_pct, seed=None):
    """
    For each clause, take: max(n, ceil(total * m_pct / 100)), capped at total.
    light=max(1,25%), standard=max(2,50%), deep=max(3,75%), full=100%.
    """
    import random
    from collections import OrderedDict
    rng = random.Random(seed)
    grouped = OrderedDict()
    for row in master_rows:
        grouped.setdefault(row["clause_base"], []).append(row)
    filtered = []
    for base, rows in grouped.items():
        total  = len(rows)
        by_pct = math.ceil(total * m_pct / 100)
        take   = min(total, max(n, by_pct))
        filtered.extend(sorted(rng.sample(rows, take), key=lambda r: rows.index(r)))
    return filtered

# ── Step 8: Write Word document ──────────────────────────
def create_docx(rows, output_path, doc_title="", sheet_label="MASTER AUDIT CHECKLIST"):
    print(f"\n[5/5] Writing: {output_path}")
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        sys.exit("ERROR: pip install python-docx")
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.75)
        sec.left_margin = sec.right_margin = Inches(0.65)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sheet_label); r.bold = True; r.font.size = Pt(13)
    if doc_title:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(doc_title); r2.italic = True; r2.font.size = Pt(9)
    doc.add_paragraph()
    HEADERS    = ["SL\nNO.", "Clause #", "Sub\nPoint", "DESCRIPTION",
                  "ORG. PROCEDURE #", "ADEQUACY\nYES / NO", "COMPLIANT\nYES / NO", "REMARKS"]
    COL_WIDTHS = [0.32, 1.30, 0.45, 2.90, 1.10, 0.72, 0.72, 1.00]
    table = doc.add_table(rows=1, cols=len(HEADERS)); table.style = "Table Grid"
    def shade(cell, color):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color); tcPr.append(shd)
    def pad(cell):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcM = OxmlElement("w:tcMar")
        for side, val in [("top",50),("bottom",50),("left",80),("right",80)]:
            n = OxmlElement(f"w:{side}"); n.set(qn("w:w"), str(val))
            n.set(qn("w:type"), "dxa"); tcM.append(n)
        tcPr.append(tcM)
    for cell, w, h in zip(table.rows[0].cells, COL_WIDTHS, HEADERS):
        cell.width = Inches(w); pad(cell)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.bold = True; run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); shade(cell, "1F4E79")
    for idx, item in enumerate(rows, start=1):
        row = table.add_row(); bg = "EBF3FB" if idx % 2 == 0 else "FFFFFF"
        vals = [str(idx), item.get("clause_base", item.get("ref","")),
                item.get("subpoint",""), item["question"], "", "", "", ""]
        for i, (cell, val) in enumerate(zip(row.cells, vals)):
            cell.width = Inches(COL_WIDTHS[i]); pad(cell); shade(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 3 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val); run.font.size = Pt(8)
    doc.save(output_path)
    print(f"   Saved -> {output_path}  ({len(rows)} items)")

# ── Step 9: Build master rows ────────────────────────────
def build_master(llm, clauses, max_items=0, progress_cb=None,
                 prompt_template=None, focus=""):
    subset = clauses[:max_items] if max_items else clauses
    total  = len(subset); master = []
    for i, clause in enumerate(subset, 1):
        if progress_cb: progress_cb(i, total, clause["ref"])
        checkpoints = extract_checkpoints(llm, clause,
                                          prompt_template=prompt_template, focus=focus)
        if not checkpoints:
            print(f"   [{i:>3}/{total}] {clause['ref']} — no checkpoints extracted"); continue
        n = len(checkpoints)
        for j, q in enumerate(checkpoints, 1):
            master.append({"clause_base": clause["ref"], "subpoint": f"{j}/{n}",
                           "ref": f"{clause['ref']} ({j}/{n})", "question": q})
        print(f"   [{i:>3}/{total}] {clause['ref']} — {n} checkpoints")
    return master

# ── Main ─────────────────────────────────────────────────
def main():
    pages              = read_pdf(PDF_PATH)
    llm                = load_model(MODEL_PATH)
    structure          = detect_structure(llm, pages)
    clause_re, sect_re = compile_patterns(structure)
    clauses            = extract_clauses(pages, clause_re, sect_re)
    if not clauses: sys.exit("\nERROR: No clauses found.")
    doc_name = os.path.splitext(os.path.basename(PDF_PATH))[0].replace("_", " ")
    print(f"\n[5/5] Generating checkpoints (Turbo Mode)...")
    master   = build_master_turbo(llm, clauses, max_items=MAX_ITEMS, batch_size=5)
    if not master: sys.exit("\nERROR: No checkpoints extracted.")
    create_docx(master, OUTPUT_DOCX, doc_title=doc_name, sheet_label="MASTER AUDIT CHECKLIST")
    print(f"\nMaster: {len(master)} checkpoints -> {OUTPUT_DOCX}")
    filtered      = filter_checkpoints(master, n=2, m_pct=50)
    filtered_path = OUTPUT_DOCX.replace(".docx", "_filtered.docx")
    create_docx(filtered, filtered_path, doc_title=doc_name,
                sheet_label="AUDIT SHEET — STANDARD DEPTH (N=2, M=50%)")
    print(f"Filtered: {len(filtered)} checkpoints -> {filtered_path}")

if __name__ == "__main__":
    main()
