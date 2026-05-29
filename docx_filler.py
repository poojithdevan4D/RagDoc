import os
import docx

def fill_quarterly_report(template_path, monthly_datas, output_path):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Quarterly template not found at {template_path}")
        
    doc = docx.Document(template_path)
    
    # Determine quarter months
    months = [d.get("month", "") for d in monthly_datas if d.get("month")]
    quarter_text = ", ".join(filter(None, months))
    
    # Replace dates/quarter header placeholders in paragraphs
    for p in doc.paragraphs:
        if "FOR QUARTER ENDING" in p.text:
            p.text = f"FOR QUARTER ENDING: {quarter_text.upper()}"
            
    # Gather and deduplicate all records from monthly reports
    all_major_projects = []
    all_indigenous_projects = []
    all_work_progress = []
    all_audits = []
    all_meetings = []
    all_defects = []
    all_contributions = []
    
    seen_major = set()
    seen_indigenous = set()
    seen_work = set()
    seen_audits = set()
    seen_meetings = set()
    seen_defects = set()
    seen_contribs = set()
    
    for m_data in monthly_datas:
        # Projects
        for proj in m_data.get("major_projects", []):
            pkey = (proj.get("item", "").strip().lower(), proj.get("part_no", "").strip().lower())
            if pkey not in seen_major:
                seen_major.add(pkey)
                all_major_projects.append(proj)
                
        # Indigenous projects
        for ip in m_data.get("indigenous_projects", []):
            if ip and ip.strip().lower() not in seen_indigenous:
                seen_indigenous.add(ip.strip().lower())
                all_indigenous_projects.append(ip.strip())
                
        # Work progress / Milestones
        for wp in m_data.get("work_progress", []):
            bullet = wp if isinstance(wp, str) else wp.get("bullet", "")
            if bullet and bullet.strip().lower() not in seen_work:
                seen_work.add(bullet.strip().lower())
                all_work_progress.append(bullet.strip())
                
        # Audits
        for audit in m_data.get("audits", []):
            akey = (audit.get("details", "").strip().lower(), audit.get("observations", "").strip().lower())
            if akey not in seen_audits:
                seen_audits.add(akey)
                all_audits.append(audit)
                
        # Meetings
        for mtg in m_data.get("meetings", []):
            desc = mtg if isinstance(mtg, str) else mtg.get("description", "")
            if desc and desc.strip().lower() not in seen_meetings:
                seen_meetings.add(desc.strip().lower())
                all_meetings.append(desc.strip())
                
        # Defects
        for defect in m_data.get("defects", []):
            dkey = (defect.get("project", "").strip().lower(), defect.get("snag", "").strip().lower())
            if dkey not in seen_defects:
                seen_defects.add(dkey)
                all_defects.append(defect)
                
        # QA Contributions
        for contrib in m_data.get("contributions", []):
            if contrib and contrib.strip().lower() not in seen_contribs:
                seen_contribs.add(contrib.strip().lower())
                all_contributions.append(contrib.strip())
                
    # Differentiate completed vs in-progress projects
    completed_projects = []
    in_progress_projects = []
    
    # Map work progress milestones to completed status
    for wp_item in all_work_progress:
        # Check if it fits as a completed project
        completed_projects.append({
            "item": wp_item,
            "part_no": "-",
            "project": "Development Progress",
            "status": "Cleared / Completed"
        })
        
    for proj in all_major_projects:
        status_lower = proj.get("status", "").lower()
        if "completed" in status_lower or "cleared" in status_lower or "passed" in status_lower:
            completed_projects.append(proj)
        else:
            in_progress_projects.append(proj)
            
    # Iterate through tables and populate them
    for table in doc.tables:
        if not table.rows:
            continue
            
        # Parse table cells to check for headers or specific rows
        header_text = "".join([cell.text for cell in table.rows[0].cells]).lower()
        
        # 1. Table 1.1 (a): STORE DEVELOPED IN FIGURES
        if "till last qtr of current" in header_text:
            if len(table.rows) > 1:
                # Write completed projects count into column index 1 ("THIS QTR ENDING")
                # and column index 2 ("TOTAL TILL DATE")
                table.rows[1].cells[1].text = str(len(completed_projects))
                table.rows[1].cells[2].text = str(len(completed_projects))
                
        # 2. Table 1.1 (b): DESCRIPTION (Stores Developed)
        elif "store description" in header_text and "type approval" in header_text:
            for idx, proj in enumerate(completed_projects):
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = proj.get("item", "")
                row.cells[2].text = proj.get("part_no", "")
                row.cells[3].text = proj.get("project", "")
                row.cells[4].text = proj.get("status", "")
                
        # 3. Table 1.2: STORES UNDER DEVELOPMENT
        elif "store description" in header_text and "clearance reference" in header_text:
            # Add major in-progress projects
            idx = 0
            for proj in in_progress_projects:
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = proj.get("item", "")
                row.cells[2].text = proj.get("part_no", "")
                row.cells[3].text = proj.get("project", "")
                row.cells[4].text = proj.get("status", "")
                idx += 1
            # Add indigenous Section B projects
            for ip in all_indigenous_projects:
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = ip
                row.cells[2].text = "-"
                row.cells[3].text = "LRU / Indigenous"
                row.cells[4].text = "Under Development"
                idx += 1
                
        # 4. Table 3.3: IMPORTANT OBSERVATIONS (DGAQA Improvements)
        elif "brief description of observation made" in header_text and "action initiated" in header_text:
            for idx, contrib in enumerate(all_contributions):
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = "Development Projects"
                row.cells[2].text = contrib
                row.cells[3].text = "Action Initiated"
                row.cells[4].text = "Completed / Implemented"
                
        # 5. Table 3.5.1 & Table 3.5.2: QUALITY AUDITS
        elif "brief detail of the audit" in header_text:
            external_audits = []
            internal_audits = []
            for audit in all_audits:
                details_lower = audit.get("details", "").lower() + audit.get("name", "").lower()
                if "external" in details_lower or "main contractor" in details_lower:
                    external_audits.append(audit)
                else:
                    internal_audits.append(audit)
                    
            if "type of audit" in header_text:
                # Table 3.5.1: Quality Audit of Main Contractor
                for idx, audit in enumerate(external_audits):
                    row = table.add_row()
                    row.cells[0].text = str(idx + 1)
                    row.cells[1].text = f"{audit.get('name', '')} - {audit.get('details', '')}"
                    row.cells[2].text = "Process/QMS Audit"
                    row.cells[3].text = audit.get("observations", "")
            else:
                # Table 3.5.2: System Quality Audit of Sub-contractor
                for idx, audit in enumerate(internal_audits):
                    row = table.add_row()
                    row.cells[0].text = str(idx + 1)
                    row.cells[1].text = f"{audit.get('name', '')} - {audit.get('details', '')}"
                    row.cells[2].text = audit.get("observations", "")
                    
        # 6. Table 3.5.3: MISCELLANEOUS ACTIVITIES (IN FIGURES)
        elif "inspector" in "".join([cell.text for row in table.rows for cell in row.cells]).lower():
            # Check row text names
            inspector_approvals = 0
            vendor_approvals = 0
            for audit in all_audits:
                details_lower = audit.get("details", "").lower() + audit.get("name", "").lower()
                if "inspector" in details_lower or "process" in details_lower:
                    inspector_approvals += 1
            for ip in all_indigenous_projects:
                if "vendor" in ip.lower():
                    vendor_approvals += 1
            
            for row in table.rows:
                row_label = row.cells[1].text.lower()
                if "inspector" in row_label:
                    row.cells[3].text = str(inspector_approvals) # THIS QTR ENDING
                elif "vendor" in row_label:
                    row.cells[3].text = str(vendor_approvals)
                    
        # 7. Table 3.5.5: DEFECT INVESTIGATIONS
        elif "accident/ incident/defect" in header_text or "brief description of accident" in header_text:
            for idx, defect in enumerate(all_defects):
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = defect.get("project", "")
                row.cells[2].text = f"{defect.get('item', '')}: {defect.get('snag', '')}"
                row.cells[3].text = defect.get("decision", "")
                
        # 8. Table 3.5.6: SUMMARY OF DI / PWR / IDI DISPOSED OF
        elif "di" in "".join([cell.text for row in table.rows for cell in row.cells]).lower() and "pwr" in "".join([cell.text for row in table.rows for cell in row.cells]).lower():
            total_defects = len(all_defects)
            for row in table.rows:
                if len(row.cells) > 2:
                    row_label = row.cells[1].text.lower().strip()
                    if row_label == "di" or row_label == "idi":
                        row.cells[2].text = str(total_defects)
                        
        # 9. Table 3.5.8: MEETINGS ATTENDED
        elif "description of the meeting" in header_text:
            for idx, mtg in enumerate(all_meetings):
                row = table.add_row()
                row.cells[0].text = str(idx + 1)
                row.cells[1].text = mtg
                row.cells[2].text = "--"
                row.cells[3].text = "Attended"
                
        # 10. Check for "Note for 3.3" or "Three major technical highlight" inside the table cells
        for row in table.rows:
            if len(row.cells) > 1:
                if "note for 3.3" in row.cells[0].text.lower() or "three major technical highlight" in row.cells[1].text.lower():
                    # Populate cell 1 with the technical highlights narrative
                    highlights_text = "Three major technical highlights / contributions for the quarter:\n\n"
                    for i, contrib in enumerate(all_contributions[:3]):
                        highlights_text += f"{i+1}. {contrib}\n\n"
                    row.cells[1].text = highlights_text
                    
    doc.save(output_path)
    print(f"Aggregated all monthly data and generated quarterly report at: {output_path}")
    return output_path
