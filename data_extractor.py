import os
import json
import pdfplumber
import docx

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ""
    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text += p.text + "\n"
    # Extract tables
    for i, table in enumerate(doc.tables):
        text += f"\n[Table {i+1}]\n"
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            # De-duplicate adjacent identical cells (due to merged cells)
            clean_row = []
            for item in row_text:
                if not clean_row or clean_row[-1] != item:
                    clean_row.append(item)
            text += " | ".join(clean_row) + "\n"
    return text

def parse_monthly_report(llm, filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        text = extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    prompt = (
        "You are an AI assistant. Analyze the monthly progress report text below and extract "
        "every single piece of information into a clean JSON format. Do not omit any project, "
        "audit, meeting, defect, or contribution.\n\n"
        "Extract the following JSON structure:\n"
        "{\n"
        "  \"month\": \"Name of Month (e.g., January)\",\n"
        "  \"major_projects\": [\n"
        "    { \"sl_no\": \"...\", \"item\": \"...\", \"part_no\": \"...\", \"project\": \"...\", \"status\": \"...\" }\n"
        "  ],\n"
        "  \"indigenous_projects\": [ \"...\" ],\n"
        "  \"work_progress\": [\n"
        "    { \"bullet\": \"...\" }\n"
        "  ],\n"
        "  \"audits\": [\n"
        "    { \"sl_no\": \"...\", \"name\": \"...\", \"details\": \"...\", \"observations\": \"...\", \"remarks\": \"...\" }\n"
        "  ],\n"
        "  \"meetings\": [ \"...\" ],\n"
        "  \"defects\": [\n"
        "    { \"sl_no\": \"...\", \"project\": \"...\", \"item\": \"...\", \"snag\": \"...\", \"decision\": \"...\" }\n"
        "  ],\n"
        "  \"contributions\": [ \"...\" ]\n"
        "}\n"
        "Only output the JSON object, do not write other text or markdown formatting.\n\n"
        "REPORT TEXT:\n"
        "--------------------\n"
        f"{text[:6000]}\n"
        "--------------------\n\n"
        "JSON OUTPUT:\n"
    )
    
    # We call the local LLM directly using the module function
    from document_summarizer import _run_local_llm
    raw_output = _run_local_llm(llm, prompt, max_tokens=1500, temperature=0.1)
    
    # Clean output if the model wrapped it in markdown
    raw_output = raw_output.strip()
    if raw_output.startswith("```json"):
        raw_output = raw_output.split("```json", 1)[1]
    if raw_output.endswith("```"):
        raw_output = raw_output.rsplit("```", 1)[0]
    raw_output = raw_output.strip()
    
    # 1. Extract only the first complete JSON block (discards extra replies/garbage)
    def extract_first_json_block(text):
        first_open = text.find('{')
        if first_open == -1:
            return text
        nest = 0
        for idx in range(first_open, len(text)):
            char = text[idx]
            if char == '{':
                nest += 1
            elif char == '}':
                nest -= 1
                if nest == 0:
                    return text[first_open:idx+1]
        return text[first_open:]
        
    cleaned = extract_first_json_block(raw_output)
    
    import re
    # 2. Inject missing commas between adjacent JSON objects or elements (e.g. ] "key" or } "key")
    cleaned = re.sub(r'\}\s*"\s*([a-zA-Z_]+)"', r'}, "\1"', cleaned)
    cleaned = re.sub(r'\]\s*"\s*([a-zA-Z_]+)"', r'], "\1"', cleaned)
    
    # Remove trailing commas inside lists/objects (invalid in Python json.loads)
    cleaned = re.sub(r',\s*([\]}])', r'\1', cleaned)
    
    try:
        return json.loads(cleaned)
    except Exception as e:
        # Try simple repair for truncated JSON by counting open/close symbols
        open_braces = cleaned.count('{')
        close_braces = cleaned.count('}')
        open_brackets = cleaned.count('[')
        close_brackets = cleaned.count(']')
        
        repaired = cleaned
        if open_brackets > close_brackets:
            repaired += ']' * (open_brackets - close_brackets)
        if open_braces > close_braces:
            repaired += '}' * (open_braces - close_braces)
            
        try:
            return json.loads(repaired)
        except Exception:
            print(f"Failed to parse LLM JSON output for {filepath}: {e}")
            print("Raw output was:", raw_output)
            print("Repaired attempted was:", repaired)
            
            # Regex extraction fallback for Month name
            month_match = re.search(r'"month"\s*:\s*"([^"]+)"', raw_output)
            month = month_match.group(1) if month_match else "Unknown"
            return {
                "month": month,
                "major_projects": [],
                "indigenous_projects": [],
                "work_progress": [],
                "audits": [],
                "meetings": [],
                "defects": [],
                "contributions": []
            }
